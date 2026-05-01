# ============================================================
# AI DOCUMENT INTELLIGENCE SYSTEM (FINAL WORKING VERSION)
# LOGIN + MULTI-FILE + RAG + CHAT UI (FREE MODEL)
# ============================================================

import tempfile
import streamlit as st
import json
import os
from dotenv import load_dotenv

st.set_page_config(page_title="AI Document Intelligence System", layout="wide")
load_dotenv()

# ------------------------------
# IMPORTS
# ------------------------------
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.documents import Document
from transformers import pipeline
from langchain_community.llms import HuggingFacePipeline

# ------------------------------
# USER STORAGE
# ------------------------------
USER_FILE = "users.json"

def load_users():
    if not os.path.exists(USER_FILE):
        data = {"users": {}, "count": 0}
        with open(USER_FILE, "w") as f:
            json.dump(data, f)
        return data

    with open(USER_FILE, "r") as f:
        try:
            data = json.load(f)
        except:
            data = {"users": {}, "count": 0}

    data.setdefault("users", {})
    data.setdefault("count", 0)
    return data

def save_users(data):
    with open(USER_FILE, "w") as f:
        json.dump(data, f)

# ------------------------------
# SESSION STATE
# ------------------------------
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

if "username" not in st.session_state:
    st.session_state.username = ""

if "user_count" not in st.session_state:
    st.session_state.user_count = load_users()["count"]

if "retriever" not in st.session_state:
    st.session_state.retriever = None

if "messages" not in st.session_state:
    st.session_state.messages = []

if "last_file" not in st.session_state:
    st.session_state.last_file = None

if "embeddings" not in st.session_state:
    st.session_state.embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2"
    )

# ------------------------------
# AUTH PAGE
# ------------------------------
def auth_page():
    st.title("🔐 Authentication")

    choice = st.radio("Select Option", ["Login", "Create Account"])
    data = load_users()

    if choice == "Login":
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")

        if st.button("Login"):
            if username in data["users"] and data["users"][username] == password:
                st.session_state.logged_in = True
                st.session_state.username = username
                data["count"] += 1
                save_users(data)
                st.session_state.user_count = data["count"]
                st.rerun()
            else:
                st.error("Invalid credentials")

    if choice == "Create Account":
        new_user = st.text_input("New Username")
        new_pass = st.text_input("New Password", type="password")

        if st.button("Create Account"):
            if new_user in data["users"]:
                st.warning("Username exists")
            elif new_user == "" or new_pass == "":
                st.warning("Fill all fields")
            else:
                data["users"][new_user] = new_pass
                save_users(data)
                st.success("Account created. Please login.")

# ------------------------------
# PROTECT APP
# ------------------------------
if not st.session_state.logged_in:
    auth_page()
    st.stop()

# ------------------------------
# MAIN UI
# ------------------------------
st.title("📄 AI Document Intelligence System")

st.sidebar.success(f"👤 {st.session_state.username}")
st.sidebar.success(f"👥 Users: {st.session_state.user_count}")

if st.sidebar.button("Logout"):
    st.session_state.logged_in = False
    st.session_state.username = ""
    st.rerun()

# ------------------------------
# FILE UPLOAD
# ------------------------------
uploaded_file = st.file_uploader(
    "Upload your file",
    type=["pdf", "txt", "docx", "xlsx"],
    accept_multiple_files=True
)

if uploaded_file:
    if st.session_state.last_file != uploaded_file.name:
        st.session_state.retriever = None
        st.session_state.messages = []
        st.session_state.last_file = uploaded_file.name

# ------------------------------
# PROCESS FILE
# ------------------------------
if uploaded_files and st.session_state.retriever is None:

    all_docs = []

    with st.spinner("Processing files..."):

        for uploaded_file in uploaded_files:

            file_type = uploaded_file.name.split(".")[-1].lower()

            if file_type not in ["pdf", "txt", "docx", "xlsx"]:
                continue

            if file_type == "pdf":
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                    tmp.write(uploaded_file.read())
                    pdf_path = tmp.name

                loader = PyPDFLoader(pdf_path)
                docs = loader.load()

            elif file_type == "txt":
                text = uploaded_file.read().decode("utf-8")
                docs = [Document(page_content=text)]

            elif file_type == "docx":
                from docx import Document as DocxDocument
                doc = DocxDocument(uploaded_file)
                text = "\n".join([p.text for p in doc.paragraphs])
                docs = [Document(page_content=text)]

            elif file_type == "xlsx":
                import pandas as pd
                df = pd.read_excel(uploaded_file)
                text = df.to_string()
                docs = [Document(page_content=text)]

            # ✅ add metadata
            for doc in docs:
                doc.metadata["source"] = uploaded_file.name

            all_docs.extend(docs)

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=100
        )

        chunks = splitter.split_documents(all_docs)

        vectorstore = FAISS.from_documents(chunks, st.session_state.embeddings)
        st.session_state.retriever = vectorstore.as_retriever(search_kwargs={"k": 2})

        st.success("Files processed successfully!")
# ------------------------------
# LLM (FREE MODEL)
# ------------------------------
from transformers import pipeline
from langchain_community.llms import HuggingFacePipeline

pipe = pipeline(
    "text-generation",   # ✅ correct task
    model="gpt2",
    max_new_tokens=150
)

llm = HuggingFacePipeline(pipeline=pipe)

# ------------------------------
# PROMPT
# ------------------------------
prompt = ChatPromptTemplate.from_messages([
    ("system", "Answer only using the given context."),
    ("human", "{question}\n\nContext:\n{context}")
])

# ------------------------------
# CHAT UI
# ------------------------------
st.divider()
st.subheader("💬 Chat")

for msg in st.session_state.messages:
    role = "user" if isinstance(msg, HumanMessage) else "assistant"
    with st.chat_message(role):
        st.write(msg.content)

# ------------------------------
# CHAT INPUT
# ------------------------------
if st.session_state.retriever:
    question = st.chat_input("Ask your question...")

    if question:
        with st.chat_message("user"):
            st.write(question)

        with st.spinner("Thinking..."):
            docs = st.session_state.retriever.invoke(question)
            context = "\n\n".join([doc.page_content for doc in docs])[:1500]

            response = llm.invoke(
                prompt.format_messages(question=question, context=context)
            )

            answer = response  # HuggingFace returns string

        with st.chat_message("assistant"):
            st.write(answer)

        st.session_state.messages.append(HumanMessage(content=question))
        st.session_state.messages.append(AIMessage(content=answer))
